from typing import List, Optional, Dict, Any, Union
import os
import uuid
from sqlalchemy.orm import Session
from sqlalchemy import or_, and_
from fastapi import UploadFile, HTTPException
from app.models.document import Document, DocumentCreate, DocumentUpdate, DocumentStatus
from app.services.vector_store import VectorStoreService
import requests
from io import BytesIO
import csv
import json
import base64
import logging
from app.models.task import Task
from datetime import datetime


class DocumentService:
    """
    文档服务
    
    注意：Document模型已更新以匹配数据库变更，从使用document_id字段改为使用id作为主键。
    数据库修改包括：
    - 删除了document_id、content、extracted_text、vector_id、doc_metadata、deleted字段
    - 新增了task_id、processing_status、bucket_name、object_key、content_type、file_size、etag、
      vector_store_id、vector_collection_name、vector_count字段
    - 将metadata字段替代了原来的doc_metadata字段
    
    所有接口从使用document_id字符串改为使用id整数作为标识符。
    """
    
    def __init__(self, db: Session, vector_store: VectorStoreService):
        self.db = db
        self.vector_store = vector_store
        
    def create_document(
            self,
            user_id: int,
            name: str,
            file_type: str,
            content: Optional[str] = None,  # 保留参数但内部处理不直接存储
            extracted_text: Optional[str] = None,  # 保留参数但内部处理不直接存储
            doc_metadata: Optional[Dict[str, Any]] = None,
            public: bool = False,
            bucket_name: Optional[str] = None,
            object_key: Optional[str] = None,
            content_type: Optional[str] = None,
            file_size: Optional[int] = None,
            etag: Optional[str] = None,
            task_id: Optional[str] = None,
            **kwargs,
        ) -> Document:
        """
        创建文档
        """
        logger = logging.getLogger(__name__)
        logger.info(f"创建文档: {name}, 类型: {file_type}, 用户ID: {user_id}")
        
        # 确保doc_metadata是字典类型
        if doc_metadata is not None and not isinstance(doc_metadata, dict):
            logger.warning(f"doc_metadata不是字典类型，将尝试转换: {type(doc_metadata)}")
            try:
                doc_metadata = dict(doc_metadata)
            except Exception as e:
                logger.error(f"转换doc_metadata为字典失败，将使用空字典: {str(e)}")
                doc_metadata = {}
        
        # 创建文档记录 - 不再设置id字段，让数据库自动生成
        document = Document(
            user_id=user_id,
            name=name,
            file_type=file_type,
            task_id=task_id,
            processing_status=DocumentStatus.PENDING,
            bucket_name=bucket_name,
            object_key=object_key,
            content_type=content_type,
            file_size=file_size,
            etag=etag,
            doc_metadata=doc_metadata or {},  # 直接设置doc_metadata字段
            **kwargs,
        )
        
        # 保存文档至数据库
        try:
            db = self.db
            db.add(document)
            db.commit()
            db.refresh(document)
            logger.info(f"文档创建成功: ID={document.id}")
            return document
        except Exception as e:
            db.rollback()
            logger.error(f"文档创建失败: {str(e)}")
            raise e
    
    def get_document(self, doc_id: int, user_id: int) -> Optional[Document]:
        """
        获取文档
        
        Args:
            doc_id: 文档ID（整数）
            user_id: 用户ID
            
        Returns:
            文档对象，如果不存在则返回None
        """
        return self.db.query(Document).filter(
            Document.id == doc_id,
            Document.user_id == user_id
        ).first()
    
    def get_documents(self, 
                     user_id: int, 
                     skip: int = 0, 
                     limit: int = 100, 
                     search: Optional[str] = None,
                     filters: Optional[Dict[str, Any]] = None) -> tuple[List[Document], int]:
        """
        获取文档列表
        
        Args:
            user_id: 用户ID
            skip: 分页起始位置
            limit: 分页大小
            search: 搜索关键词
            filters: 其他过滤条件，key为字段名，value为过滤值
            
        Returns:
            tuple: (文档列表, 总数)
        """
        query = self.db.query(Document).filter(
            Document.user_id == user_id
        )
        
        # 关键词搜索
        if search:
            query = query.filter(
                Document.name.ilike(f"%{search}%")
            )
        
        # 应用其他过滤条件
        if filters:
            for field, value in filters.items():
                if hasattr(Document, field):
                    query = query.filter(getattr(Document, field) == value)
        
        # 获取总数
        total = query.count()
        
        # 分页并排序
        documents = query.order_by(Document.created_at.desc()).offset(skip).limit(limit).all()
        
        return documents, total
    
    def get_documents_with_tasks(self, 
                           user_id: int, 
                           skip: int = 0, 
                           limit: int = 100, 
                           search: Optional[str] = None,
                           filters: Optional[Dict[str, Any]] = None) -> tuple[List[Dict[str, Any]], int]:
        """
        获取文档列表，包含最新任务状态
        
        Args:
            user_id: 用户ID
            skip: 分页起始位置
            limit: 分页大小
            search: 搜索关键词
            filters: 其他过滤条件，key为字段名，value为过滤值
            
        Returns:
            tuple: (文档列表(包含任务), 总数)
        """
        # 获取文档列表
        documents, total = self.get_documents(user_id, skip, limit, search, filters)
        
        # 为每个文档添加最新任务信息
        results = []
        for doc in documents:
            # 查询最新任务
            latest_task = self.db.query(Task).filter(
                Task.document_id == doc.id
            ).order_by(Task.created_at.desc()).first()
            
            # 构建响应对象
            doc_dict = doc.__dict__.copy()
            doc_dict.pop('_sa_instance_state', None)
            doc_dict['latest_task'] = latest_task.__dict__.copy() if latest_task else None
            if doc_dict['latest_task']:
                doc_dict['latest_task'].pop('_sa_instance_state', None)
                
            results.append(doc_dict)
            
        return results, total
    
    def update_document(self, 
                       doc_id: int, 
                       user_id: int, 
                       document_update: DocumentUpdate) -> Optional[Document]:
        """更新文档"""
        logger = logging.getLogger(__name__)
        
        db_document = self.get_document(doc_id, user_id)
        if not db_document:
            return None
        
        update_data = document_update.dict(exclude_unset=True)
        
        # 处理metadata字段
        if "metadata" in update_data:
            metadata = update_data["metadata"]
            # 确保metadata是字典类型
            if metadata is not None and not isinstance(metadata, dict):
                logger.warning(f"更新元数据不是字典类型，将转换为字典: {type(metadata)}")
                try:
                    metadata = dict(metadata)
                except Exception as e:
                    logger.error(f"转换更新元数据为字典失败，将使用空字典: {str(e)}")
                    metadata = {}
            # 直接使用doc_metadata字段（不再通过元素访问器）
            update_data["doc_metadata"] = metadata
            del update_data["metadata"]
        
        # 移除不再存在的字段
        for field in ["content", "extracted_text", "vector_id", "deleted"]:
            if field in update_data:
                logger.warning(f"{field}字段已不再直接存储于document模型中，将被忽略")
                del update_data[field]
        
        # 处理提取文本用于向量化（如果提供）
        extracted_text = document_update.extracted_text if hasattr(document_update, 'extracted_text') else None
        if extracted_text and hasattr(self, 'vector_store'):
            # 限制文本长度
            max_length = 2048
            if len(extracted_text) > max_length:
                logger.warning(f"更新的提取文本超过{max_length}字符，将被截断")
                text_for_vector = extracted_text[:max_length]
            else:
                text_for_vector = extracted_text
            
            # 构建向量元数据
            vector_metadata = {
                "id": doc_id,  # 使用id而非document_id
                "name": db_document.name,
                "file_type": db_document.file_type
            }
            
            # 合并其他元数据
            if db_document.doc_metadata:
                vector_metadata.update(db_document.doc_metadata)
            
            # 添加新向量
            if db_document.vector_store_id:
                # 删除旧向量
                self.vector_store.delete_texts([db_document.vector_store_id])
            
            # 添加新向量
            vector_ids = self.vector_store.add_texts(
                texts=[text_for_vector],
                metadatas=[vector_metadata]
            )
            
            if vector_ids:
                # 更新向量存储ID
                update_data["vector_store_id"] = vector_ids[0]
                # 添加向量数量信息
                update_data["vector_count"] = 1
        
        # 更新文档记录
        for key, value in update_data.items():
            setattr(db_document, key, value)
        
        self.db.commit()
        self.db.refresh(db_document)
        return db_document
    
    def delete_document(self, doc_id: int, user_id: int, permanent: bool = True) -> bool:
        """
        删除文档
        
        Args:
            doc_id: 文档ID
            user_id: 用户ID（用于验证权限）
            permanent: 是否永久删除，False则标记为已删除状态
            
        Returns:
            删除是否成功
        """
        logger = logging.getLogger(__name__)
        logger.info(f"删除文档: ID={doc_id}, 用户ID={user_id}, 永久删除={permanent}")
        
        document = self.db.query(Document).filter(Document.id == doc_id, Document.user_id == user_id).first()
        if not document:
            logger.warning(f"找不到要删除的文档: ID={doc_id}, 用户ID={user_id}")
            return False
        
        if permanent:
            logger.info(f"永久删除文档: ID={doc_id}")
            # TODO: 删除对象存储中的文件
            self.db.delete(document)
        else:
            logger.info(f"标记文档为已删除: ID={doc_id}")
            # 仅将状态更新为已删除
            db_document = self.db.query(Document).filter(Document.id == doc_id).first()
            db_document.processing_status = DocumentStatus.DELETED
            # 可以在metadata中添加删除时间等信息
            if db_document.doc_metadata is None:
                db_document.doc_metadata = {}
            db_document.doc_metadata["deleted_at"] = datetime.utcnow().isoformat()
        
        self.db.commit()
        logger.info(f"文档删除操作完成: ID={doc_id}")
        return True
    
    async def process_file(self, 
                    file: UploadFile, 
                    user_id: int,
                    doc_metadata: Optional[Dict[str, Any]] = None) -> Document:
        """处理上传文件"""
        logger = logging.getLogger(__name__)
        
        # 记录开始处理文件
        logger.info(f"开始处理文件: {file.filename}, 内容类型: {file.content_type}")
        
        # 确保doc_metadata是字典类型
        if doc_metadata is not None and not isinstance(doc_metadata, dict):
            logger.warning(f"process_file中doc_metadata不是字典类型，将尝试转换: {type(doc_metadata)}")
            try:
                doc_metadata = dict(doc_metadata)
            except Exception as e:
                logger.error(f"转换doc_metadata为字典失败，将使用空字典: {str(e)}")
                doc_metadata = {}
        
        # 获取文件内容
        content = await file.read()
        
        # 记录文件大小
        file_size = len(content)
        logger.info(f"文件大小: {file_size} 字节")
        
        # 提取文本
        extracted_text = await self._extract_text_from_file(file.filename, file.content_type, content)
        
        # 获取文件类型
        file_extension = os.path.splitext(file.filename)[1].lower() if file.filename else ""
        file_type = file_extension.lstrip(".") or file.content_type or "unknown"
        logger.info(f"文件类型识别为: {file_type}")
        
        # 判断是否为二进制文件
        binary_types = ['pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx', 'zip', 'rar']
        is_binary = file_type in binary_types
        
        # 如果是二进制文件，使用Base64编码
        if is_binary:
            logger.info(f"检测到二进制文件类型: {file_type}，将进行Base64编码")
            try:
                # 检查是否包含NUL字符
                has_null = b'\x00' in content
                if has_null:
                    logger.info("文件包含NUL字符，确认为二进制文件")
                
                # 对内容进行Base64编码，添加前缀以便后续识别
                content_to_save = "__BASE64__" + base64.b64encode(content).decode('ascii')
                logger.info(f"Base64编码完成，编码后大小: {len(content_to_save)} 字节")
            except Exception as e:
                logger.error(f"Base64编码失败: {str(e)}")
                # 出错时尝试忽略错误字符的普通解码
                content_to_save = content.decode('utf-8', errors='ignore')
        else:
            # 文本文件正常处理
            logger.info(f"文本文件类型: {file_type}，使用普通UTF-8解码")
            content_to_save = content.decode('utf-8', errors='ignore')
        
        # 创建元数据
        file_metadata = {
            "filename": file.filename,
            "content_type": file.content_type,
            "size": file_size,
            "is_binary_encoded": is_binary,
            **(doc_metadata or {})
        }

        # 生成S3相关信息
        bucket_name = "documents"
        object_key = f"{user_id}/{uuid.uuid4()}/{file.filename}"
        etag = str(uuid.uuid4())  # 模拟ETag
        content_type = file.content_type or "application/octet-stream"
        
        # 创建文档
        logger.info("准备创建文档记录")
        return self.create_document(
            user_id=user_id,
            name=file.filename or "Unnamed document",
            file_type=file_type,
            content=content_to_save,
            extracted_text=extracted_text,
            doc_metadata=file_metadata,
            bucket_name=bucket_name,
            object_key=object_key,
            content_type=content_type,
            file_size=file_size,
            etag=etag
        )
    
    async def _extract_text_from_file(self, 
                               filename: Optional[str], 
                               content_type: Optional[str], 
                               content: bytes) -> str:
        """从文件中提取文本"""
        # 根据文件类型调用不同的提取方法
        if not filename:
            return content.decode('utf-8', errors='ignore')
            
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext in ['.txt']:
            # 纯文本文件
            return content.decode('utf-8', errors='ignore')
            
        elif file_ext in ['.pdf']:
            # PDF文件
            try:
                import PyPDF2
                from io import BytesIO
                
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except Exception as e:
                print(f"PDF提取文本错误: {str(e)}")
                return content.decode('utf-8', errors='ignore')
                
        elif file_ext in ['.doc', '.docx']:
            # Word文档
            try:
                import docx
                from io import BytesIO
                
                doc = docx.Document(BytesIO(content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except Exception as e:
                print(f"Word提取文本错误: {str(e)}")
                return content.decode('utf-8', errors='ignore')

        elif file_ext in ['.xls', '.xlsx']:
            # Excel文件
            try:
                import openpyxl
                from io import BytesIO
                
                workbook = openpyxl.load_workbook(BytesIO(content))
                text = ""
                for sheet in workbook:
                    text += f"Sheet: {sheet.title}\n"
                    for row in sheet.iter_rows(values_only=True):
                        text += "\t".join([str(cell) if cell is not None else "" for cell in row]) + "\n"
                return text
            except Exception as e:
                print(f"Excel提取文本错误: {str(e)}")
                return content.decode('utf-8', errors='ignore')
                
        elif file_ext in ['.csv']:
            # CSV文件
            try:
                csv_text = content.decode('utf-8', errors='ignore')
                reader = csv.reader(csv_text.splitlines())
                
                extracted_text = ""
                headers = next(reader, [])
                extracted_text += "Headers: " + ", ".join(headers) + "\n\n"
                
                for i, row in enumerate(reader):
                    if i < 1000:  # 限制行数，避免过大
                        extracted_text += "Row " + str(i+1) + ": " + ", ".join(row) + "\n"
                    else:
                        extracted_text += f"\n... (showing first 1000 rows out of more)"
                        break
                        
                return extracted_text
            except Exception as e:
                print(f"CSV提取文本错误: {str(e)}")
                return content.decode('utf-8', errors='ignore')

        elif file_ext in ['.json']:
            # JSON文件
            try:
                json_text = content.decode('utf-8', errors='ignore')
                json_data = json.loads(json_text)
                # 美化输出
                pretty_json = json.dumps(json_data, indent=2, ensure_ascii=False)
                return pretty_json
            except Exception as e:
                print(f"JSON提取文本错误: {str(e)}")
                return content.decode('utf-8', errors='ignore')
                
        elif file_ext in ['.md']:
            # Markdown文件
            try:
                md_text = content.decode('utf-8', errors='ignore')
                # 可以考虑将Markdown转换为纯文本，但简单起见，直接返回原始内容
                return md_text
            except Exception as e:
                print(f"Markdown提取文本错误: {str(e)}")
                return content.decode('utf-8', errors='ignore')
                
        elif file_ext in ['.htm', '.html']:
            # HTML文件
            try:
                from bs4 import BeautifulSoup
                
                html_text = content.decode('utf-8', errors='ignore')
                soup = BeautifulSoup(html_text, 'html.parser')
                
                # 移除脚本和样式
                for script in soup(["script", "style"]):
                    script.extract()
                
                # 获取文本
                text = soup.get_text(separator='\n')
                # 处理多余空白行
                lines = (line.strip() for line in text.splitlines())
                chunks = (line for line in lines if line)
                text = '\n'.join(chunks)
                return text
            except Exception as e:
                print(f"HTML提取文本错误: {str(e)}")
                return content.decode('utf-8', errors='ignore')
        
        # 默认尝试解码为文本
        return content.decode('utf-8', errors='ignore')
        
    async def load_from_web(self, url: str, user_id: int, metadata: Optional[Dict[str, Any]] = None) -> Document:
        """从网页加载文档"""
        logger = logging.getLogger(__name__)
        logger.info(f"从网页加载文档: {url}")
        
        try:
            # 发送请求获取网页内容
            response = requests.get(url, timeout=10, allow_redirects=True)
            response.raise_for_status()  # 如果状态码不是200，抛出异常
            
            # 获取内容类型
            content_type = response.headers.get('Content-Type', '')
            logger.info(f"网页内容类型: {content_type}")
            
            # 解析内容类型
            is_html = 'text/html' in content_type.lower()
            is_json = 'application/json' in content_type.lower()
            is_text = 'text/' in content_type.lower()
            
            # 获取内容
            content = response.content
            file_size = len(content)
            logger.info(f"网页内容大小: {file_size} 字节")
            
            # 提取网页标题
            title = url
            if is_html:
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(content, 'html.parser')
                    if soup.title and soup.title.string:
                        title = soup.title.string.strip()
                except Exception as e:
                    logger.error(f"提取网页标题失败: {str(e)}")
            
            # 提取文本
            if is_html:
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(content, 'html.parser')
                    
                    # 移除脚本和样式
                    for script in soup(["script", "style"]):
                        script.extract()
                    
                    # 获取文本
                    text = soup.get_text(separator='\n')
                    # 处理多余空白行
                    lines = (line.strip() for line in text.splitlines())
                    chunks = (line for line in lines if line)
                    extracted_text = '\n'.join(chunks)
                except Exception as e:
                    logger.error(f"提取HTML文本失败: {str(e)}")
                    extracted_text = content.decode('utf-8', errors='ignore')
            elif is_json:
                try:
                    json_data = json.loads(content.decode('utf-8', errors='ignore'))
                    extracted_text = json.dumps(json_data, indent=2, ensure_ascii=False)
                except Exception as e:
                    logger.error(f"解析JSON失败: {str(e)}")
                    extracted_text = content.decode('utf-8', errors='ignore')
            else:
                # 其他类型直接解码
                extracted_text = content.decode('utf-8', errors='ignore')
            
            # 生成文档元数据
            doc_metadata = {
                "url": url,
                "content_type": content_type,
                "size": file_size,
                **(metadata or {})
            }
            
            # 文件类型推断
            if is_html:
                file_type = "html"
            elif is_json:
                file_type = "json"
            elif 'text/plain' in content_type.lower():
                file_type = "txt"
            elif 'text/csv' in content_type.lower():
                file_type = "csv"
            else:
                # 尝试从URL获取文件扩展名
                from urllib.parse import urlparse
                parsed_url = urlparse(url)
                path = parsed_url.path
                file_ext = os.path.splitext(path)[1].lower()
                if file_ext:
                    file_type = file_ext.lstrip('.')
                else:
                    file_type = "html" if is_html else "txt"  # 默认
            
            # 文件内容处理
            content_str = content.decode('utf-8', errors='ignore')
            
            # 创建文档
            logger.info(f"从网页创建文档: {title}, 类型: {file_type}")
            return self.create_document(
                user_id=user_id,
                name=title or url,
                file_type=file_type,
                content=content_str,
                extracted_text=extracted_text,
                doc_metadata=doc_metadata,
                content_type=content_type,
                file_size=file_size
            )
            
        except Exception as e:
            logger.error(f"从网页加载文档失败: {str(e)}")
            raise e
            
    async def load_from_directory(self, directory_path: str, user_id: int, recursive: bool = True) -> List[Document]:
        """从目录加载文档"""
        logger = logging.getLogger(__name__)
        logger.info(f"从目录加载文档: {directory_path}, 递归={recursive}")
        
        # 检查目录是否存在
        if not os.path.exists(directory_path):
            logger.error(f"目录不存在: {directory_path}")
            raise HTTPException(status_code=400, detail=f"目录不存在: {directory_path}")
        
        if not os.path.isdir(directory_path):
            logger.error(f"路径不是目录: {directory_path}")
            raise HTTPException(status_code=400, detail=f"路径不是目录: {directory_path}")
        
        # 支持的文件类型
        supported_extensions = [
            '.txt', '.pdf', '.doc', '.docx', '.xls', '.xlsx', 
            '.csv', '.json', '.md', '.html', '.htm'
        ]
        
        documents = []
        
        # 列出文件
        for root, dirs, files in os.walk(directory_path):
            # 如果不递归，只处理顶层目录
            if not recursive and root != directory_path:
                continue
                
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                # 跳过不支持的文件类型
                if file_ext not in supported_extensions:
                    logger.info(f"跳过不支持的文件类型: {file_path}")
                    continue
                
                logger.info(f"处理文件: {file_path}")
                
                try:
                    # 读取文件内容
                    with open(file_path, 'rb') as f:
                        content = f.read()
                    
                    # 获取文件大小
                    file_size = len(content)
                    
                    # 根据文件扩展名确定mimetype
                    mimetypes = {
                        '.txt': 'text/plain',
                        '.pdf': 'application/pdf',
                        '.doc': 'application/msword',
                        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        '.xls': 'application/vnd.ms-excel',
                        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        '.csv': 'text/csv',
                        '.json': 'application/json',
                        '.md': 'text/markdown',
                        '.html': 'text/html',
                        '.htm': 'text/html'
                    }
                    content_type = mimetypes.get(file_ext, 'application/octet-stream')
                    
                    # 提取文本
                    extracted_text = await self._extract_text_from_file(file, content_type, content)
                    
                    # 处理内容（二进制文件使用Base64编码）
                    binary_types = ['.pdf', '.doc', '.docx', '.xls', '.xlsx']
                    is_binary = file_ext in binary_types
                    
                    if is_binary:
                        # Base64编码
                        try:
                            content_to_save = "__BASE64__" + base64.b64encode(content).decode('ascii')
                        except Exception as e:
                            logger.error(f"Base64编码失败: {str(e)}")
                            content_to_save = content.decode('utf-8', errors='ignore')
                    else:
                        # 文本文件
                        content_to_save = content.decode('utf-8', errors='ignore')
                    
                    # 创建元数据
                    file_metadata = {
                        "filename": file,
                        "path": os.path.relpath(file_path, directory_path),
                        "content_type": content_type,
                        "size": file_size,
                        "is_binary_encoded": is_binary
                    }
                    
                    # 生成S3相关信息
                    bucket_name = "documents"
                    object_key = f"{user_id}/{uuid.uuid4()}/{file}"
                    etag = str(uuid.uuid4())  # 模拟ETag
                    
                    # 创建文档
                    document = self.create_document(
                        user_id=user_id,
                        name=file,
                        file_type=file_ext.lstrip("."),
                        content=content_to_save,
                        extracted_text=extracted_text,
                        doc_metadata=file_metadata,
                        bucket_name=bucket_name,
                        object_key=object_key,
                        content_type=content_type,
                        file_size=file_size,
                        etag=etag
                    )
                    
                    documents.append(document)
                    logger.info(f"成功添加文档: {file}")
                    
                except Exception as e:
                    logger.error(f"处理文件 {file_path} 失败: {str(e)}")
                    # 继续处理下一个文件
                    continue
        
        return documents
    
    async def create_custom_document(self, 
                              user_id: int, 
                              name: str, 
                              content: str, 
                              file_type: str = "txt",
                              metadata: Optional[Dict[str, Any]] = None) -> Document:
        """创建自定义文档（手动输入）"""
        logger = logging.getLogger(__name__)
        logger.info(f"创建自定义文档: {name}, 类型: {file_type}")
        
        # 内容和大小
        content_bytes = content.encode('utf-8')
        file_size = len(content_bytes)
        
        # 根据文件类型确定mimetype
        content_type_map = {
            "txt": "text/plain",
            "md": "text/markdown",
            "json": "application/json",
            "html": "text/html",
            "csv": "text/csv"
        }
        content_type = content_type_map.get(file_type, "text/plain")
        
        # 创建元数据
        doc_metadata = {
            "source": "manual",
            "content_type": content_type,
            "size": file_size,
            **(metadata or {})
        }
        
        # 生成S3相关信息
        bucket_name = "documents"
        object_key = f"{user_id}/{uuid.uuid4()}/{name}.{file_type}"
        etag = str(uuid.uuid4())  # 模拟ETag
        
        # 创建文档
        return self.create_document(
            user_id=user_id,
            name=name,
            file_type=file_type,
            content=content,
            extracted_text=content,  # 直接使用内容作为提取文本
            doc_metadata=doc_metadata,
            bucket_name=bucket_name,
            object_key=object_key,
            content_type=content_type,
            file_size=file_size,
            etag=etag
        )
    
    def update_document_status(self, doc_id: int, status: Union[str, DocumentStatus], message: Optional[str] = None) -> Optional[Document]:
        """
        更新文档处理状态
        
        Args:
            doc_id: 文档ID
            status: 处理状态，可以是 DocumentStatus 枚举或字符串
            message: 状态消息
            
        Returns:
            更新后的文档对象，如果文档不存在则返回None
        """
        logger = logging.getLogger(__name__)
        logger.info(f"更新文档状态: ID={doc_id}, 状态={status}")
        
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            logger.warning(f"找不到文档: ID={doc_id}")
            return None
        
        # 处理 status 参数，确保它是字符串类型
        if isinstance(status, DocumentStatus):
            document.processing_status = status.value
        else:
            document.processing_status = status
        
        # 如果传入了消息，可以保存到metadata中
        if message and document.doc_metadata is not None:
            if isinstance(document.doc_metadata, dict):
                document.doc_metadata["status_message"] = message
        
        self.db.commit()
        self.db.refresh(document)
        logger.info(f"文档状态已更新: ID={doc_id}, 状态={document.processing_status}")
        return document
    
    async def validate_file(self, doc_id: int, file_path: str) -> Dict[str, Any]:
        """
        验证文件
        
        Args:
            doc_id: 文档ID
            file_path: 文件路径
            
        Returns:
            Dict[str, Any]: 包含验证结果的字典
        """
        logger = logging.getLogger(__name__)
        logger.info(f"验证文件: doc_id={doc_id}, file_path={file_path}")
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 获取文件大小
        file_size = os.path.getsize(file_path)
        
        # 检查文件类型
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 支持的文件类型
        supported_extensions = ['.txt', '.pdf', '.doc', '.docx', '.xls', '.xlsx', 
                               '.csv', '.json', '.md', '.html', '.htm']
        
        if file_ext not in supported_extensions:
            logger.warning(f"不支持的文件类型: {file_ext}")
        
        # 检查文件大小限制
        max_size = 50 * 1024 * 1024  # 50MB
        if file_size > max_size:
            logger.warning(f"文件过大: {file_size} 字节")
        
        return {
            "validated": True,
            "file_size": file_size,
            "file_type": file_ext.lstrip('.'),
            "file_path": file_path
        }
    
    async def preprocess_text(self, doc_id: int, file_path: str, extracted_text: Optional[str] = None, **kwargs) -> Dict[str, Any]:
        """
        预处理文本
        
        Args:
            doc_id: 文档ID
            file_path: 文件路径
            extracted_text: 已提取的文本（可选）
            **kwargs: 额外的关键字参数
            
        Returns:
            Dict[str, Any]: 包含预处理结果的字典
        """
        logger = logging.getLogger(__name__)
        logger.info(f"预处理文本: doc_id={doc_id}")
        
        # 记录额外接收到的参数
        if kwargs:
            logger.debug(f"preprocess_text 收到额外参数: {kwargs}")
            
        # 如果没有提供文本，则尝试读取文件
        if not extracted_text:
            logger.info("未提供文本，尝试读取文件")
            if not os.path.exists(file_path):
                logger.error(f"文件不存在: {file_path}")
                raise FileNotFoundError(f"文件不存在: {file_path}")
                
            # 获取文件扩展名
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 读取文件内容
            with open(file_path, 'rb') as f:
                content = f.read()
                
            # 提取文本
            extracted_text = await self._extract_text_from_file(file_path, None, content)
        
        # 预处理文本
        processed_text = extracted_text
        
        # 1. 去除多余空格
        processed_text = " ".join(processed_text.split())
        
        # 2. 标准化换行符
        processed_text = processed_text.replace("\r\n", "\n").replace("\r", "\n")
        
        # 3. 去除特殊字符
        # processed_text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', '', processed_text)
        
        return {
            "processed_text": processed_text,
            "original_length": len(extracted_text) if extracted_text else 0,
            "processed_length": len(processed_text)
        }
    
    async def vectorize_document(self, doc_id: int, file_path: str, processed_text: Optional[str] = None, **kwargs) -> Dict[str, Any]:
        """
        向量化文档
        
        Args:
            doc_id: 文档ID
            file_path: 文件路径
            processed_text: 预处理后的文本
            **kwargs: 额外的关键字参数
            
        Returns:
            Dict[str, Any]: 包含向量化结果的字典
        """
        logger = logging.getLogger(__name__)
        logger.info(f"向量化文档: doc_id={doc_id}")
        
        # 记录额外接收到的参数
        if kwargs:
            logger.debug(f"vectorize_document 收到额外参数: {kwargs}")
        
        # 获取文档信息
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            logger.error(f"找不到文档: {doc_id}")
            raise ValueError(f"找不到文档: {doc_id}")
        
        # 如果没有提供处理过的文本，则尝试使用传入的文件路径读取文件
        if not processed_text and os.path.exists(file_path):
            logger.info("未提供处理过的文本，尝试从文件提取")
            # 读取文件内容
            with open(file_path, 'rb') as f:
                content = f.read()
                
            # 提取文本
            extracted_text = await self._extract_text_from_file(file_path, None, content)
            
            # 预处理文本
            preprocess_result = await self.preprocess_text(doc_id, file_path, extracted_text)
            processed_text = preprocess_result["processed_text"]
        
        if not processed_text:
            logger.error("无法获取文本进行向量化")
            raise ValueError("无法获取文本进行向量化")
        
        # 限制文本长度
        max_length = 2048
        if len(processed_text) > max_length:
            logger.warning(f"文本超过{max_length}字符，将被截断")
            text_for_vector = processed_text[:max_length]
        else:
            text_for_vector = processed_text
        
        # 构建向量元数据
        vector_metadata = {
            "id": doc_id,  # 使用id而非document_id
            "name": document.name,
            "file_type": document.file_type
        }
        
        # 合并其他元数据
        if document.doc_metadata:
            vector_metadata.update(document.doc_metadata)
        
        # 向量化处理 - 实际会调用向量服务，这里简化处理
        embeddings = None
        if hasattr(self, 'vector_store'):
            try:
                # 调用向量服务生成嵌入
                vectors = self.vector_store.generate_embeddings([text_for_vector])
                if vectors and len(vectors) > 0:
                    embeddings = vectors[0]
            except Exception as e:
                logger.error(f"生成向量嵌入失败: {str(e)}")
        
        return {
            "text_length": len(text_for_vector),
            "embeddings": embeddings,
            "metadata": vector_metadata,
            "processed_text": text_for_vector  # 返回用于向量化的文本
        }
    
    async def store_document_vectors(self, doc_id: int, file_path: str, 
                                    processed_text: Optional[str] = None, 
                                    embeddings: Optional[List[float]] = None,
                                    **kwargs) -> Dict[str, Any]:
        """
        存储文档向量
        
        Args:
            doc_id: 文档ID
            file_path: 文件路径
            processed_text: 预处理后的文本
            embeddings: 嵌入向量（可选）
            **kwargs: 额外的关键字参数
            
        Returns:
            Dict[str, Any]: 包含存储结果的字典
        """
        logger = logging.getLogger(__name__)
        logger.info(f"存储文档向量: doc_id={doc_id}")
        
        # 记录额外接收到的参数
        if kwargs:
            logger.debug(f"store_document_vectors 收到额外参数: {kwargs}")
        
        # 获取文档信息
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            logger.error(f"找不到文档: {doc_id}")
            raise ValueError(f"找不到文档: {doc_id}")
        
        # 如果没有提供向量，则生成向量
        if not embeddings or not processed_text:
            logger.info("未提供向量或处理过的文本，尝试生成")
            vectorize_result = await self.vectorize_document(doc_id, file_path, processed_text)
            embeddings = vectorize_result.get("embeddings")
            processed_text = vectorize_result.get("processed_text")
            metadata = vectorize_result.get("metadata")
        else:
            # 构建向量元数据
            metadata = {
                "id": doc_id,
                "name": document.name,
                "file_type": document.file_type
            }
            
            # 合并其他元数据
            if document.doc_metadata:
                metadata.update(document.doc_metadata)
        
        # 如果没有embeddings，则无法存储
        if not embeddings or not processed_text:
            logger.error("无法获取向量进行存储")
            raise ValueError("无法获取向量进行存储")
        
        # 存储向量 - 实际会调用向量服务，这里简化处理
        vector_ids = None
        if hasattr(self, 'vector_store'):
            try:
                # 如果文档已有向量ID，则先删除
                if document.vector_store_id:
                    self.vector_store.delete_texts([document.vector_store_id])
                
                # 存储新向量
                vector_ids = self.vector_store.add_texts(
                    texts=[processed_text],
                    metadatas=[metadata]
                )
                
                # 更新文档的向量ID和数量
                if vector_ids:
                    document.vector_store_id = vector_ids[0]
                    document.vector_count = 1
                    document.processing_status = DocumentStatus.COMPLETED
                    self.db.commit()
                    self.db.refresh(document)
            except Exception as e:
                logger.error(f"存储向量失败: {str(e)}")
                raise e
        
        return {
            "vector_ids": vector_ids,
            "vector_count": 1 if vector_ids else 0,
            "status": "success" if vector_ids else "failed"
        }
    
    def update_document_content(self, doc_id: int, content: str) -> Optional[Document]:
        """
        更新文档文本内容
        
        由于文档模型已不再包含content字段，这里需要将内容保存在其他地方
        例如保存到对象存储或更新metadata
        
        Args:
            doc_id: 文档ID
            content: 文档内容
            
        Returns:
            更新后的文档对象，如果文档不存在则返回None
        """
        logger = logging.getLogger(__name__)
        logger.info(f"更新文档内容: ID={doc_id}")
        
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            logger.warning(f"找不到文档: ID={doc_id}")
            return None
        
        # 将内容保存到对象存储或者更新metadata中的内容标记
        # 这里需要根据实际情况实现
        # 例如，可以将内容保存到对象存储中
        if document.bucket_name and document.object_key:
            # 实现保存内容到对象存储的逻辑
            pass
        
        # 或者将内容保存在metadata中（不推荐用于大型文档）
        if document.doc_metadata is None:
            document.doc_metadata = {}
        if isinstance(document.doc_metadata, dict):
            document.doc_metadata["has_extracted_content"] = True
        
        self.db.commit()
        self.db.refresh(document)
        logger.info(f"文档内容已更新: ID={doc_id}")
        return document
    
    async def extract_text_from_file_path(self, doc_id: int, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        从文件路径中提取文本的适配器方法
        
        这个方法适配Celery任务，接受文件路径参数，然后调用内部的_extract_text_from_file方法
        
        Args:
            doc_id: 文档ID
            file_path: 文件路径
            **kwargs: 额外的关键字参数，包括来自任务管道的参数
            
        Returns:
            Dict[str, Any]: 包含提取结果的字典
        """
        logger = logging.getLogger(__name__)
        logger.info(f"从文件提取文本: doc_id={doc_id}, file_path={file_path}")
        
        # 记录额外接收到的参数
        if kwargs:
            logger.debug(f"extract_text_from_file_path 收到额外参数: {kwargs}")
        
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 读取文件内容
        with open(file_path, 'rb') as f:
            content = f.read()
        
        # 提取文件名
        filename = os.path.basename(file_path)
        
        # 调用内部方法提取文本
        extracted_text = await self._extract_text_from_file(filename, None, content)
        
        return {
            "extracted_text": extracted_text,
            "text_length": len(extracted_text) if extracted_text else 0
        } 